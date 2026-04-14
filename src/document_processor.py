"""
Document Processor - Handles PDF, DOCX, and TXT files
"""

import os
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from datetime import datetime

from utils.text_splitter import SmartTextSplitter
from src.config import Config


@dataclass
class Document:
    """Represents a processed document"""
    doc_id: str
    filename: str
    file_type: str
    content: str
    chunks: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    file_size: int = 0
    chunk_count: int = 0


class DocumentProcessor:
    """Handles multi-format document ingestion and processing"""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".txt", ".md"}

    def __init__(self):
        self.splitter = SmartTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP
        )
        self.processed_docs: Dict[str, Document] = {}

    def process_file(self, file_path: str) -> Document:
        """Process a single file and return a Document object"""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.SUPPORTED_FORMATS}")

        print(f"📄 Processing: {path.name}")

        # Extract text based on file type
        if ext == ".pdf":
            content = self._extract_pdf(file_path)
        elif ext == ".docx":
            content = self._extract_docx(file_path)
        elif ext in (".txt", ".md"):
            content = self._extract_text(file_path)

        # Generate unique doc ID
        doc_id = self._generate_doc_id(file_path, content)

        # Split into chunks
        chunks = self.splitter.split_text(content)

        # Add chunk metadata
        enriched_chunks = []
        for i, chunk in enumerate(chunks):
            # chunk is a dict with 'text', 'chunk_index', etc. from TextSplitter
            chunk_text = chunk["text"] if isinstance(chunk, dict) else str(chunk)
            enriched_chunks.append({
                "chunk_id": f"{doc_id}_chunk_{i}",
                "text": chunk_text,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "source": path.name,
                "doc_id": doc_id
            })

        # Build document
        doc = Document(
            doc_id=doc_id,
            filename=path.name,
            file_type=ext.lstrip("."),
            content=content,
            chunks=enriched_chunks,
            metadata={
                "source": str(path),
                "file_type": ext,
                "word_count": len(content.split()),
                "char_count": len(content),
            },
            file_size=path.stat().st_size,
            chunk_count=len(chunks)
        )

        self.processed_docs[doc_id] = doc
        print(f"   ✅ Extracted {len(content.split())} words → {len(chunks)} chunks")
        return doc

    def process_directory(self, dir_path: str) -> List[Document]:
        """Process all supported files in a directory"""
        path = Path(dir_path)
        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")

        documents = []
        files = [f for f in path.iterdir() if f.suffix.lower() in self.SUPPORTED_FORMATS]

        print(f"\n📁 Processing {len(files)} files from {dir_path}")

        for file_path in files:
            try:
                doc = self.process_file(str(file_path))
                documents.append(doc)
            except Exception as e:
                print(f"   ⚠️ Skipped {file_path.name}: {e}")

        print(f"\n✅ Successfully processed {len(documents)}/{len(files)} files")
        return documents

    def process_text(self, text: str, source_name: str = "manual_input") -> Document:
        """Process raw text directly"""
        doc_id = self._generate_doc_id(source_name, text)
        chunks = self.splitter.split_text(text)

        enriched_chunks = []
        for i, chunk in enumerate(chunks):
            # chunk is a dict with 'text', 'chunk_index', etc. from TextSplitter
            chunk_text = chunk["text"] if isinstance(chunk, dict) else str(chunk)
            enriched_chunks.append({
                "chunk_id": f"{doc_id}_chunk_{i}",
                "text": chunk_text,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "source": source_name,
                "doc_id": doc_id
            })

        doc = Document(
            doc_id=doc_id,
            filename=source_name,
            file_type="text",
            content=text,
            chunks=enriched_chunks,
            metadata={"source": source_name, "word_count": len(text.split())},
            chunk_count=len(chunks)
        )
        self.processed_docs[doc_id] = doc
        return doc

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{text}")
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

    def _extract_text(self, file_path: str) -> str:
        """Extract text from TXT/MD files"""
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Could not decode file: {file_path}")

    def _generate_doc_id(self, file_path: str, content: str) -> str:
        """Generate a unique document ID"""
        hash_input = f"{file_path}{content[:100]}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:12]

    def get_stats(self) -> Dict[str, Any]:
        """Return processing statistics"""
        return {
            "total_documents": len(self.processed_docs),
            "total_chunks": sum(d.chunk_count for d in self.processed_docs.values()),
            "total_words": sum(d.metadata.get("word_count", 0) for d in self.processed_docs.values()),
            "file_types": list({d.file_type for d in self.processed_docs.values()})
        }